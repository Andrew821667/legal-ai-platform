# -*- coding: utf-8 -*-
import json
from io import BytesIO

from docx import Document

from src.services.stage2_document_generator import Stage2DocumentGenerator


def test_generate_corrected_docx():
    generator = Stage2DocumentGenerator()
    source_doc = Document()
    source_doc.add_paragraph("Исходный текст договора")
    source_buffer = BytesIO()
    source_doc.save(source_buffer)

    recommendations = [
        {
            "section_number": "5",
            "section_title": "Ответственность",
            "original_text": "Сторона несет неограниченную ответственность",
            "proposed_text": "Ответственность ограничивается суммой договора",
            "reason": "Снижение финансового риска",
            "action_type": "modify",
            "priority": "critical",
            "source": "section_analysis",
        }
    ]

    result_bytes = generator.generate_corrected_docx(
        base_docx_bytes=source_buffer.getvalue(),
        accepted_recommendations=recommendations,
        source_file_name="draft.docx",
    )

    assert isinstance(result_bytes, bytes)
    assert len(result_bytes) > 200

    result_doc = Document(BytesIO(result_bytes))
    full_text = "\n".join([p.text for p in result_doc.paragraphs if p.text])
    assert "Внесенные правки" in full_text
    assert "Ответственность ограничивается суммой договора" in full_text


def test_generate_protocol_docx_and_json():
    generator = Stage2DocumentGenerator()
    recommendations = [
        {
            "section_number": "2",
            "section_title": "Оплата",
            "original_text": "Оплата 90 дней",
            "proposed_text": "Оплата 30 дней",
            "reason": "Снижение дебиторской нагрузки",
            "action_type": "modify",
            "priority": "important",
            "source": "template_comparison",
        }
    ]

    protocol_docx = generator.generate_disagreement_protocol_docx(
        accepted_recommendations=recommendations,
        source_file_name="draft.docx",
    )
    assert isinstance(protocol_docx, bytes)
    assert len(protocol_docx) > 200

    protocol_json = generator.generate_disagreement_protocol_json(recommendations)
    payload = json.loads(protocol_json)
    assert len(payload) == 1
    assert payload[0]["section"].startswith("2.")
    assert payload[0]["after"] == "Оплата 30 дней"


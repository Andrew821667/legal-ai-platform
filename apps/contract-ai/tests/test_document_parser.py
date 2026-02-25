"""
Тестирование Document Parser
"""
import sys
import os
sys.path.insert(0, os.getcwd())

from src.services.document_parser import DocumentParser
from docx import Document
from docx.shared import Pt
from pathlib import Path


def create_test_docx():
    """Создаёт тестовый DOCX для проверки парсера"""
    print("Создание тестового DOCX...")

    doc = Document()

    # Заголовок
    doc.add_heading('ДОГОВОР ПОСТАВКИ №123', 0)

    # Преамбула
    doc.add_paragraph(
        'ООО "Поставщик", ИНН 1234567890, именуемое в дальнейшем "Поставщик", '
        'в лице Генерального директора Иванова И.И., действующего на основании Устава, '
        'с одной стороны, и ООО "Покупатель", ИНН 0987654321, именуемое в дальнейшем '
        '"Покупатель", в лице директора Петрова П.П., действующего на основании Устава, '
        'с другой стороны, заключили настоящий договор о нижеследующем:'
    )

    # 1. Предмет договора
    doc.add_heading('1. Предмет договора', 1)
    doc.add_paragraph(
        '1.1. Поставщик обязуется поставить, а Покупатель принять и оплатить товар '
        'в количестве и ассортименте согласно Спецификации (Приложение №1).'
    )
    doc.add_paragraph(
        '1.2. Наименование товара: Пшеница продовольственная, 3 класс.'
    )

    # 2. Цена и порядок расчётов
    doc.add_heading('2. Цена и порядок расчётов', 1)
    doc.add_paragraph(
        '2.1. Общая стоимость товара составляет 1 000 000 рублей (Один миллион рублей).'
    )
    doc.add_paragraph(
        '2.2. Оплата производится в следующем порядке: 30% предоплата в течение 5 банковских дней, '
        'остальные 70% в течение 10 дней после поставки.'
    )

    # 3. Сроки
    doc.add_heading('3. Сроки поставки', 1)
    doc.add_paragraph(
        '3.1. Поставка товара осуществляется до 31.12.2025.'
    )

    # 4. Ответственность сторон
    doc.add_heading('4. Ответственность сторон', 1)
    doc.add_paragraph(
        '4.1. За нарушение сроков поставки Поставщик уплачивает пени в размере 0.1% '
        'от стоимости товара за каждый день просрочки.'
    )

    # Таблица
    doc.add_heading('Приложение №1. Спецификация', 1)
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Light Grid Accent 1'

    # Заголовки
    headers = ['№', 'Наименование', 'Количество', 'Цена']
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    # Данные
    table.rows[1].cells[0].text = '1'
    table.rows[1].cells[1].text = 'Пшеница 3 класс'
    table.rows[1].cells[2].text = '100 тонн'
    table.rows[1].cells[3].text = '10 000 руб/тонна'

    table.rows[2].cells[0].text = 'ИТОГО'
    table.rows[2].cells[1].text = ''
    table.rows[2].cells[2].text = '100 тонн'
    table.rows[2].cells[3].text = '1 000 000 руб'

    # Подписи
    doc.add_paragraph('\n')
    doc.add_paragraph('Подписано: 15.10.2025')

    # Сохраняем
    test_file = 'tests/fixtures/test_contract.docx'
    os.makedirs('tests/fixtures', exist_ok=True)
    doc.save(test_file)

    print(f"   ✓ Тестовый DOCX создан: {test_file}")
    return test_file


def test_document_parser():
    """Тестирует Document Parser"""
    print("=" * 60)
    print("ТЕСТИРОВАНИЕ DOCUMENT PARSER")
    print("=" * 60)

    # Тест 1: Создание парсера
    print("\n1. Инициализация Document Parser...")
    try:
        parser = DocumentParser()
        print(f"   ✓ Parser создан")
        print(f"   - Поддерживаемые форматы: {parser.supported_formats}")
    except Exception as e:
        print(f"   ✗ Ошибка: {e}")
        return False

    # Тест 2: Создание тестового файла
    print("\n2. Создание тестового договора...")
    try:
        test_file = create_test_docx()
    except Exception as e:
        print(f"   ✗ Ошибка создания файла: {e}")
        return False

    # Тест 3: Парсинг DOCX
    print("\n3. Парсинг DOCX файла...")
    try:
        xml_result = parser.parse(test_file)
        print(f"   ✓ Парсинг выполнен")
        print(f"   - Размер XML: {len(xml_result)} символов")

        # Проверяем наличие ключевых элементов
        if '<contract>' in xml_result:
            print("   ✓ Корневой элемент <contract> найден")
        if '<metadata>' in xml_result:
            print("   ✓ Секция <metadata> найдена")
        if '<parties>' in xml_result:
            print("   ✓ Секция <parties> найдена")
        if '<clauses>' in xml_result:
            print("   ✓ Секция <clauses> найдена")
        if '<tables>' in xml_result:
            print("   ✓ Секция <tables> найдена")

    except Exception as e:
        print(f"   ✗ Ошибка парсинга: {e}")
        import traceback
        traceback.print_exc()
        return False

    # Тест 4: Валидация XML
    print("\n4. Валидация XML...")
    try:
        is_valid = parser.validate_xml(xml_result)
        if is_valid:
            print("   ✓ XML валиден")
        else:
            print("   ✗ XML не валиден")
            return False
    except Exception as e:
        print(f"   ✗ Ошибка валидации: {e}")
        return False

    # Тест 5: Сохранение результата
    print("\n5. Сохранение результата...")
    try:
        output_file = 'tests/fixtures/test_contract.xml'
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(xml_result)
        print(f"   ✓ XML сохранён: {output_file}")
    except Exception as e:
        print(f"   ✗ Ошибка сохранения: {e}")
        return False

    # Тест 6: Извлечение tracked changes
    print("\n6. Тест извлечения tracked changes...")
    try:
        changes = parser.extract_tracked_changes(test_file)
        print(f"   ✓ Извлечено изменений: {len(changes)}")
        if len(changes) == 0:
            print("   ℹ В тестовом документе нет tracked changes (это нормально)")
    except Exception as e:
        print(f"   ✗ Ошибка: {e}")
        return False

    # Тест 7: Вывод части XML
    print("\n7. Превью XML (первые 1000 символов):")
    print("-" * 60)
    print(xml_result[:1000])
    print("-" * 60)

    print("\n" + "=" * 60)
    print("✓ ВСЕ ТЕСТЫ ПРОЙДЕНЫ")
    print("=" * 60)
    print(f"\nФайлы созданы:")
    print(f"  - Тестовый DOCX: {test_file}")
    print(f"  - Результат XML: tests/fixtures/test_contract.xml")

    return True


if __name__ == "__main__":
    success = test_document_parser()
    sys.exit(0 if success else 1)

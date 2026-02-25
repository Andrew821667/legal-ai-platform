# -*- coding: utf-8 -*-
"""
Quick Export Agent - Fast export of contracts to various formats
"""
import os
from typing import Dict, Any, Optional
from datetime import datetime
from loguru import logger

from .base_agent import BaseAgent, AgentResult
from ..services.llm_gateway import LLMGateway
from ..models.database import Contract, ExportLog


class QuickExportAgent(BaseAgent):
    """
    Agent for quick export of contracts

    Capabilities:
    - Export to DOCX (with formatting)
    - Export to PDF
    - Export to plain text
    - Export metadata (JSON)
    - Batch export
    - Export logging
    """

    def __init__(
        self,
        llm_gateway: LLMGateway,
        db_session,
        config: Optional[Dict[str, Any]] = None
    ):
        super().__init__(llm_gateway, db_session, config)
        self.export_dir = self.config.get('export_dir', 'data/exports')
        os.makedirs(self.export_dir, exist_ok=True)

    def get_name(self) -> str:
        return "QuickExportAgent"

    def get_system_prompt(self) -> str:
        return "You are a document export assistant."

    def execute(self, state: Dict[str, Any]) -> AgentResult:
        """
        Execute quick export

        Expected state:
        - contract_id: ID of contract to export
        - export_format: 'docx', 'pdf', 'txt', 'json', 'all'
        - include_analysis: Include analysis results (default: False)
        - user_id: ID of user requesting export

        Returns:
        - file_paths: Dict of exported file paths
        - export_log_id: ID of export log record
        """
        try:
            contract_id = state.get('contract_id')
            export_format = state.get('export_format', 'docx')
            include_analysis = state.get('include_analysis', False)
            user_id = state.get('user_id')

            if not contract_id:
                return AgentResult(
                    success=False,
                    data={},
                    error="Missing contract_id"
                )

            logger.info(f"Quick export: contract {contract_id}, format {export_format}")

            # Get contract
            contract = self.db.query(Contract).filter(
                Contract.id == contract_id
            ).first()

            if not contract:
                return AgentResult(
                    success=False,
                    data={},
                    error=f"Contract {contract_id} not found"
                )

            # Export based on format
            file_paths = {}

            if export_format == 'all':
                formats = ['docx', 'pdf', 'txt', 'json']
            else:
                formats = [export_format]

            for fmt in formats:
                try:
                    file_path = self._export_format(contract, fmt, include_analysis)
                    file_paths[fmt] = file_path
                except Exception as e:
                    logger.error(f"Failed to export {fmt}: {e}")
                    file_paths[fmt] = None

            # Log export
            export_log = self._log_export(contract_id, export_format, file_paths, user_id)

            logger.info(f"Export complete: {len(file_paths)} files")

            return AgentResult(
                success=True,
                data={
                    'contract_id': contract_id,
                    'file_paths': file_paths,
                    'export_log_id': export_log.id if export_log else None
                },
                metadata={'message': f"Exported to {len(file_paths)} format(s)"}
            )

        except Exception as e:
            logger.error(f"Quick export failed: {e}")
            return AgentResult(
                success=False,
                data={},
                error=str(e)
            )

    def _export_format(
        self,
        contract: Contract,
        format: str,
        include_analysis: bool
    ) -> str:
        """Export contract to specific format"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_name = f"{contract.file_name.replace('.docx', '')}_{timestamp}"

        if format == 'docx':
            return self._export_docx(contract, base_name, include_analysis)
        elif format == 'pdf':
            return self._export_pdf(contract, base_name, include_analysis)
        elif format == 'txt':
            return self._export_txt(contract, base_name, include_analysis)
        elif format == 'json':
            return self._export_json(contract, base_name, include_analysis)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _export_docx(self, contract: Contract, base_name: str, include_analysis: bool) -> str:
        """Export to DOCX with proper formatting"""
        output_path = os.path.join(self.export_dir, f"{base_name}.docx")

        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Create document
            doc = Document()

            # Title
            title = doc.add_heading(f'Экспорт договора: {contract.file_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata section
            doc.add_heading('Информация о документе', 1)
            metadata_table = doc.add_table(rows=5, cols=2)
            metadata_table.style = 'Light Grid Accent 1'

            metadata_rows = [
                ('Файл:', contract.file_name),
                ('Тип документа:', contract.document_type or 'Не определён'),
                ('Тип договора:', contract.contract_type or 'Не определён'),
                ('Дата загрузки:', contract.upload_date.strftime('%d.%m.%Y %H:%M') if contract.upload_date else 'N/A'),
                ('Статус:', contract.status or 'Не определён'),
            ]

            for i, (label, value) in enumerate(metadata_rows):
                row = metadata_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = str(value)

            # Analysis results if requested
            if include_analysis and contract.analysis_results:
                doc.add_page_break()
                doc.add_heading('Результаты анализа', 1)

                for analysis in contract.analysis_results:
                    doc.add_heading(f'Анализ версии {getattr(analysis, "version", "N/A")}', 2)

                    if hasattr(analysis, 'risks'):
                        doc.add_heading('Выявленные риски', 3)

                        for risk in analysis.risks[:10]:  # Limit to 10
                            p = doc.add_paragraph(style='List Bullet')
                            severity = getattr(risk, 'severity', 'unknown').upper()
                            risk_type = getattr(risk, 'risk_type', 'general')
                            description = getattr(risk, 'description', 'Нет описания')

                            run = p.add_run(f'[{severity}] {risk_type}: ')
                            run.bold = True

                            # Color code by severity
                            if severity == 'CRITICAL':
                                run.font.color.rgb = RGBColor(220, 53, 69)  # Red
                            elif severity == 'HIGH':
                                run.font.color.rgb = RGBColor(253, 126, 20)  # Orange
                            elif severity == 'MEDIUM':
                                run.font.color.rgb = RGBColor(255, 193, 7)  # Yellow
                            else:
                                run.font.color.rgb = RGBColor(40, 167, 69)  # Green

                            p.add_run(description)

            # Footer
            doc.add_page_break()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run(
                f'Документ создан автоматически системой Contract AI\n'
                f'{datetime.now().strftime("%d.%m.%Y %H:%M")}'
            )
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)

            # Save document
            doc.save(output_path)

            logger.info(f"DOCX export completed: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            # Fallback: copy original file if exists
            import shutil
            if os.path.exists(contract.file_path):
                try:
                    shutil.copy(contract.file_path, output_path)
                    logger.info(f"DOCX export fallback: copied original file")
                    return output_path
                except:
                    pass

            # Last resort: create placeholder
            with open(output_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
                f.write(f"DOCX Export Failed: {str(e)}\n")
                f.write(f"Contract: {contract.file_name}\n")
            return output_path.replace('.docx', '.txt')

    def _export_pdf(self, contract: Contract, base_name: str, include_analysis: bool) -> str:
        """Export to PDF with professional formatting"""
        output_path = os.path.join(self.export_dir, f"{base_name}.pdf")

        try:
            from ..utils.pdf_generator import PDFGenerator

            # Prepare data for PDF
            data = {
                'summary': f'Договор: {contract.file_name}',
                'statistics': {
                    'Тип документа': contract.document_type or 'Не определён',
                    'Тип договора': contract.contract_type or 'Не определён',
                    'Дата загрузки': contract.upload_date.strftime('%d.%m.%Y %H:%M') if contract.upload_date else 'N/A',
                    'Статус': contract.status or 'Не определён',
                    'Уровень риска': contract.risk_level or 'Не оценен',
                },
                'risks': [],
                'recommendations': []
            }

            # Add analysis results if requested
            if include_analysis and contract.analysis_results:
                for analysis in contract.analysis_results:
                    if hasattr(analysis, 'risks'):
                        for risk in analysis.risks[:10]:  # Limit to 10 risks
                            data['risks'].append({
                                'severity': getattr(risk, 'severity', 'unknown'),
                                'type': getattr(risk, 'risk_type', 'general'),
                                'description': getattr(risk, 'description', ''),
                                'impact': getattr(risk, 'impact', ''),
                                'recommendation': getattr(risk, 'recommendation', '')
                            })

            # Metadata
            metadata = {
                'Файл': contract.file_name,
                'ID': str(contract.id),
                'Экспортировано': datetime.now().strftime('%d.%m.%Y %H:%M')
            }

            # Generate PDF
            generator = PDFGenerator()
            generator.generate_contract_report(
                output_path=output_path,
                title=f'Экспорт договора: {contract.file_name}',
                data=data,
                metadata=metadata
            )

            logger.info(f"PDF export completed: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            # Fallback to simple text file
            with open(output_path.replace('.pdf', '.txt'), 'w', encoding='utf-8') as f:
                f.write(f"PDF Export Failed: {str(e)}\n")
                f.write(f"Contract: {contract.file_name}\n")
            return output_path.replace('.pdf', '.txt')

    def _export_txt(self, contract: Contract, base_name: str, include_analysis: bool) -> str:
        """Export to plain text"""
        output_path = os.path.join(self.export_dir, f"{base_name}.txt")

        content = []
        content.append("=" * 80)
        content.append(f"CONTRACT: {contract.file_name}")
        content.append("=" * 80)
        content.append(f"Type: {contract.document_type}")
        content.append(f"Upload Date: {contract.upload_date}")
        content.append(f"Status: {contract.status}")
        content.append("")

        if include_analysis and contract.analysis_results:
            content.append("ANALYSIS RESULTS")
            content.append("-" * 80)
            for analysis in contract.analysis_results:
                content.append(f"Analysis ID: {analysis.id}")
                content.append(f"Version: {analysis.version}")
                content.append("")

        content.append("=" * 80)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))

        logger.info(f"TXT export: {output_path}")
        return output_path

    def _export_json(self, contract: Contract, base_name: str, include_analysis: bool) -> str:
        """Export metadata to JSON"""
        import json
        output_path = os.path.join(self.export_dir, f"{base_name}.json")

        data = {
            'id': contract.id,
            'file_name': contract.file_name,
            'file_path': contract.file_path,
            'document_type': contract.document_type,
            'contract_type': contract.contract_type,
            'upload_date': contract.upload_date.isoformat() if contract.upload_date else None,
            'status': contract.status,
            'risk_level': contract.risk_level,
            'meta_info': contract.meta_info
        }

        if include_analysis and contract.analysis_results:
            data['analysis_results'] = [
                {
                    'id': a.id,
                    'version': a.version,
                    'created_at': a.created_at.isoformat() if a.created_at else None
                }
                for a in contract.analysis_results
            ]

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        logger.info(f"JSON export: {output_path}")
        return output_path

    def _log_export(
        self,
        contract_id: str,
        export_format: str,
        file_paths: Dict[str, str],
        user_id: Optional[str]
    ) -> Optional[ExportLog]:
        """Log export to database"""
        try:
            # Find first successful export path
            file_path = None
            for path in file_paths.values():
                if path:
                    file_path = path
                    break

            if not file_path:
                return None

            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

            export_log = ExportLog(
                contract_id=contract_id,
                export_format=export_format,
                file_path=file_path,
                file_size=file_size,
                exported_by=user_id
            )

            self.db.add(export_log)
            self.db.commit()
            self.db.refresh(export_log)

            return export_log

        except Exception as e:
            logger.error(f"Failed to log export: {e}")
            return None


__all__ = ["QuickExportAgent"]

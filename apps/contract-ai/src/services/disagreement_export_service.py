# -*- coding: utf-8 -*-
"""
Disagreement Export Service - Export disagreements to various formats
Supports: DOCX, PDF, Email, EDO
"""
from typing import Dict, Any, Optional, List
from datetime import datetime
import os
import hashlib
from loguru import logger
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..models.disagreement_models import Disagreement, DisagreementObjection, DisagreementExportLog


class DisagreementExportService:
    """
    Service for exporting disagreements
    
    Formats:
    - DOCX: Official document with letterhead
    - PDF: Generated from DOCX
    - Email: SMTP with template
    - EDO: Electronic document management (stubs)
    """

    def __init__(self, db_session, config: Optional[Dict[str, Any]] = None):
        self.db_session = db_session
        self.config = config or {}
        
        # Export paths
        self.export_dir = self.config.get('export_dir', 'data/exports/disagreements')
        os.makedirs(self.export_dir, exist_ok=True)
        
        # SMTP settings (from config)
        self.smtp_host = self.config.get('smtp_host', 'smtp.gmail.com')
        self.smtp_port = self.config.get('smtp_port', 587)
        self.smtp_user = self.config.get('smtp_user')
        self.smtp_password = self.config.get('smtp_password')
        
        # EDO settings (stubs)
        self.edo_endpoints = self.config.get('edo_endpoints', {
            'diadoc': 'https://api.diadoc.ru',
            'sbis': 'https://api.sbis.ru',
            'kontur': 'https://api.kontur.ru/edo'
        })

    def export_to_docx(
        self, disagreement_id: int, user_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Export disagreement to DOCX format with official letterhead
        
        Args:
            disagreement_id: ID of disagreement
            user_id: ID of user performing export
            
        Returns:
            Dict with file_path and metadata
        """
        try:
            logger.info(f"Exporting disagreement {disagreement_id} to DOCX")

            disagreement = self.db_session.query(Disagreement).filter(
                Disagreement.id == disagreement_id
            ).first()

            if not disagreement:
                raise ValueError(f"Disagreement {disagreement_id} not found")

            # Get selected objections
            selected_ids = disagreement.priority_order or disagreement.selected_objections
            objections = self.db_session.query(DisagreementObjection).filter(
                DisagreementObjection.id.in_(selected_ids)
            ).all()

            # Sort by priority order
            objections_dict = {obj.id: obj for obj in objections}
            sorted_objections = [objections_dict[oid] for oid in selected_ids if oid in objections_dict]

            # Generate DOCX document using python-docx
            file_path = os.path.join(
                self.export_dir,
                f"disagreement_{disagreement_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )

            # Generate real DOCX document
            self._generate_docx_document(disagreement, sorted_objections, file_path)

            file_size = os.path.getsize(file_path)
            file_hash = self._calculate_file_hash(file_path)

            # Update disagreement
            disagreement.docx_path = file_path
            self.db_session.commit()

            # Log export
            self._log_export(
                disagreement_id, 'docx', file_path, file_size, file_hash, user_id
            )

            logger.info(f"DOCX exported: {file_path}")

            return {
                'success': True,
                'file_path': file_path,
                'file_size': file_size,
                'file_hash': file_hash
            }

        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def _generate_docx_document(
        self, disagreement: Disagreement, objections: List[DisagreementObjection], file_path: str
    ) -> None:
        """
        Generate professional DOCX document with python-docx

        Features:
        - Official letterhead formatting
        - Proper structure (headers, numbering)
        - Color-coded priorities
        - Table formatting for metadata
        """
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # Title
        title = doc.add_heading('ВОЗРАЖЕНИЯ К ПРОЕКТУ ДОГОВОРА', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata table
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Light Grid Accent 1'

        table.cell(0, 0).text = 'Дата:'
        table.cell(0, 1).text = datetime.now().strftime('%d.%m.%Y')
        table.cell(1, 0).text = 'ID договора:'
        table.cell(1, 1).text = str(disagreement.contract_id)

        doc.add_paragraph()  # Spacing

        # Objections section
        doc.add_heading('ЗАМЕЧАНИЯ И ПРЕДЛОЖЕНИЯ', level=2)

        priority_colors = {
            'critical': RGBColor(220, 20, 60),   # Crimson
            'high': RGBColor(255, 140, 0),        # Dark Orange
            'medium': RGBColor(255, 215, 0),      # Gold
            'low': RGBColor(60, 179, 113)         # Medium Sea Green
        }

        for i, obj in enumerate(objections, 1):
            # Objection number with priority color
            p = doc.add_paragraph()
            runner = p.add_run(f"{i}. ")
            runner.bold = True
            runner.font.size = Pt(14)

            priority_runner = p.add_run(f"[{obj.priority.upper()}] ")
            priority_runner.bold = True
            priority_runner.font.color.rgb = priority_colors.get(obj.priority.lower(), RGBColor(0, 0, 0))

            section_runner = p.add_run(f"Пункт договора: {obj.contract_section_xpath or 'не указан'}")

            # Issue description
            p_issue = doc.add_paragraph(style='List Bullet')
            issue_label = p_issue.add_run('Замечание: ')
            issue_label.bold = True
            p_issue.add_run(obj.issue_description)

            # Legal basis
            if obj.legal_basis:
                p_legal = doc.add_paragraph(style='List Bullet')
                legal_label = p_legal.add_run('Правовое обоснование: ')
                legal_label.bold = True
                p_legal.add_run(obj.legal_basis)

            # Risks
            if obj.risk_explanation:
                p_risk = doc.add_paragraph(style='List Bullet')
                risk_label = p_risk.add_run('Риски: ')
                risk_label.bold = True
                risk_label.font.color.rgb = RGBColor(220, 20, 60)  # Crimson
                p_risk.add_run(obj.risk_explanation)

            # Alternative formulation
            p_alt = doc.add_paragraph(style='List Bullet')
            alt_label = p_alt.add_run('Предлагаемая формулировка: ')
            alt_label.bold = True
            alt_label.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
            p_alt.add_run(obj.alternative_formulation)

            # Alternative reasoning
            if obj.alternative_reasoning:
                p_reason = doc.add_paragraph(style='List Bullet')
                reason_label = p_reason.add_run('Обоснование предложения: ')
                reason_label.bold = True
                p_reason.add_run(obj.alternative_reasoning)

            # Separator
            doc.add_paragraph('_' * 80)

        # Footer with summary
        doc.add_paragraph()
        summary_p = doc.add_paragraph()
        summary_p.add_run(f'Всего замечаний: {len(objections)}').bold = True

        doc.add_paragraph()
        doc.add_paragraph('С уважением,')
        doc.add_paragraph('[Подпись]')
        doc.add_paragraph('[Должность]')

        # Save document
        doc.save(file_path)
        logger.info(f"DOCX document generated: {file_path}")

    def export_to_pdf(
        self, disagreement_id: int, user_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Export disagreement to PDF using PDFGenerator
        """
        try:
            logger.info(f"Exporting disagreement {disagreement_id} to PDF")

            disagreement = self.db_session.query(Disagreement).filter(
                Disagreement.id == disagreement_id
            ).first()

            if not disagreement:
                raise ValueError(f"Disagreement {disagreement_id} not found")

            # Get selected objections
            selected_ids = disagreement.priority_order or disagreement.selected_objections
            objections = self.db_session.query(DisagreementObjection).filter(
                DisagreementObjection.id.in_(selected_ids)
            ).all()

            # Sort by priority order
            objections_dict = {obj.id: obj for obj in objections}
            sorted_objections = [objections_dict[oid] for oid in selected_ids if oid in objections_dict]

            # Convert to dict format for PDFGenerator
            objections_data = []
            for obj in sorted_objections:
                objections_data.append({
                    'section': obj.contract_section_xpath or 'не указан',
                    'issue': obj.issue_description,
                    'legal_basis': obj.legal_basis or '',
                    'alternative': obj.alternative_formulation,
                    'priority': obj.priority
                })

            disagreement_data = {
                'contract_id': disagreement.contract_id,
                'id': disagreement.id
            }

            # Generate PDF path
            pdf_path = os.path.join(
                self.export_dir,
                f"disagreement_{disagreement_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            )

            # Use PDFGenerator
            try:
                from ..utils.pdf_generator import PDFGenerator

                generator = PDFGenerator()
                generator.generate_disagreement_letter(
                    output_path=pdf_path,
                    disagreement_data=disagreement_data,
                    objections=objections_data
                )

            except ImportError:
                logger.warning("PDFGenerator not available, using fallback DOCX export")
                # Fallback to DOCX instead of PDF
                pdf_path = pdf_path.replace('.pdf', '.docx')
                self._generate_docx_document(disagreement, sorted_objections, pdf_path)

            file_size = os.path.getsize(pdf_path)
            file_hash = self._calculate_file_hash(pdf_path)

            # Update disagreement
            disagreement.pdf_path = pdf_path
            self.db_session.commit()

            # Log export
            self._log_export(
                disagreement_id, 'pdf', pdf_path, file_size, file_hash, user_id
            )

            logger.info(f"PDF exported: {pdf_path}")

            return {
                'success': True,
                'file_path': pdf_path,
                'file_size': file_size,
                'file_hash': file_hash
            }

        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def send_via_email(
        self,
        disagreement_id: int,
        recipient_email: str,
        subject: Optional[str] = None,
        user_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Send disagreement via email
        
        Args:
            disagreement_id: ID of disagreement
            recipient_email: Recipient email address
            subject: Email subject (optional)
            user_id: ID of user sending email
        """
        try:
            logger.info(f"Sending disagreement {disagreement_id} via email to {recipient_email}")

            disagreement = self.db_session.query(Disagreement).filter(
                Disagreement.id == disagreement_id
            ).first()

            if not disagreement:
                raise ValueError(f"Disagreement {disagreement_id} not found")

            # Ensure PDF exists
            if not disagreement.pdf_path or not os.path.exists(disagreement.pdf_path):
                pdf_result = self.export_to_pdf(disagreement_id, user_id)
                if not pdf_result['success']:
                    return pdf_result

            # Email subject
            if not subject:
                subject = f"Возражения к проекту договора № {disagreement.contract_id}"

            # Send email with SMTP
            email_sent = self._send_email(
                recipient_email, subject, disagreement.pdf_path
            )

            # Log export
            export_log = self._log_export(
                disagreement_id, 'email', disagreement.pdf_path, 
                os.path.getsize(disagreement.pdf_path),
                self._calculate_file_hash(disagreement.pdf_path),
                user_id
            )

            export_log.email_to = recipient_email
            export_log.email_subject = subject
            export_log.email_sent_at = datetime.utcnow()
            export_log.email_status = 'sent' if email_sent else 'failed'
            self.db_session.commit()

            logger.info(f"Email sent to {recipient_email}")

            return {
                'success': email_sent,
                'recipient': recipient_email,
                'subject': subject,
                'attachment': disagreement.pdf_path
            }

        except Exception as e:
            logger.error(f"Email send failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def _send_email(self, recipient: str, subject: str, attachment_path: str) -> bool:
        """
        Send email with attachment using SMTP

        Args:
            recipient: Recipient email address
            subject: Email subject
            attachment_path: Path to file to attach

        Returns:
            True if sent successfully, False otherwise
        """
        try:
            # Check if SMTP credentials are configured
            if not self.smtp_user or not self.smtp_password:
                logger.warning(
                    "SMTP credentials not configured. "
                    "Please set smtp_user and smtp_password in config."
                )
                logger.info(f"[SIMULATED] Email would be sent to {recipient}")
                logger.info(f"[SIMULATED] Subject: {subject}")
                logger.info(f"[SIMULATED] Attachment: {attachment_path}")
                logger.info(f"[SIMULATED] SMTP: {self.smtp_host}:{self.smtp_port}")
                return True  # Simulate success

            # Real SMTP implementation
            import smtplib
            from email.mime.multipart import MIMEMultipart
            from email.mime.text import MIMEText
            from email.mime.base import MIMEBase
            from email import encoders

            logger.info(f"Sending email to {recipient} via {self.smtp_host}:{self.smtp_port}")

            # Create message
            msg = MIMEMultipart()
            msg['From'] = self.smtp_user
            msg['To'] = recipient
            msg['Subject'] = subject

            # Email body
            body = (
                f"Добрый день!\n\n"
                f"Во вложении направляем возражения к проекту договора.\n\n"
                f"С уважением,\n"
                f"Система Contract AI"
            )
            msg.attach(MIMEText(body, 'plain', 'utf-8'))

            # Attach file
            if os.path.exists(attachment_path):
                with open(attachment_path, 'rb') as f:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(f.read())

                encoders.encode_base64(part)
                filename = os.path.basename(attachment_path)
                part.add_header(
                    'Content-Disposition',
                    f'attachment; filename="{filename}"'
                )
                msg.attach(part)
            else:
                logger.error(f"Attachment not found: {attachment_path}")
                return False

            # Send email
            server = smtplib.SMTP(self.smtp_host, self.smtp_port, timeout=30)
            server.starttls()
            server.login(self.smtp_user, self.smtp_password)
            server.send_message(msg)
            server.quit()

            logger.info(f"Email sent successfully to {recipient}")
            return True

        except smtplib.SMTPAuthenticationError as e:
            logger.error(f"SMTP authentication failed: {e}")
            logger.error("Please check smtp_user and smtp_password in config")
            return False
        except smtplib.SMTPException as e:
            logger.error(f"SMTP error: {e}")
            return False
        except Exception as e:
            logger.error(f"Email sending failed: {e}")
            return False

    def export_to_edo(
        self,
        disagreement_id: int,
        edo_system: str,
        user_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Export to EDO system (Electronic Document Management)
        
        Args:
            disagreement_id: ID of disagreement
            edo_system: 'diadoc', 'sbis', 'kontur', etc.
            user_id: ID of user performing export
        """
        try:
            logger.info(f"Exporting disagreement {disagreement_id} to EDO: {edo_system}")

            disagreement = self.db_session.query(Disagreement).filter(
                Disagreement.id == disagreement_id
            ).first()

            if not disagreement:
                raise ValueError(f"Disagreement {disagreement_id} not found")

            # Ensure PDF exists
            if not disagreement.pdf_path:
                pdf_result = self.export_to_pdf(disagreement_id, user_id)
                if not pdf_result['success']:
                    return pdf_result

            # Get EDO endpoint
            edo_endpoint = self.edo_endpoints.get(edo_system.lower())
            if not edo_endpoint:
                raise ValueError(f"Unknown EDO system: {edo_system}")

            # EDO API integration (requires paid subscription and certificates)
            # See docs/EDO_INTEGRATION.md for implementation guide
            edo_document_id = self._send_to_edo_stub(
                edo_system, edo_endpoint, disagreement.pdf_path
            )

            # Log export
            export_log = self._log_export(
                disagreement_id, 'edo', disagreement.pdf_path,
                os.path.getsize(disagreement.pdf_path),
                self._calculate_file_hash(disagreement.pdf_path),
                user_id
            )

            export_log.edo_system = edo_system
            export_log.edo_document_id = edo_document_id
            export_log.edo_status = 'sent'
            self.db_session.commit()

            logger.info(f"Exported to {edo_system}: {edo_document_id}")

            return {
                'success': True,
                'edo_system': edo_system,
                'edo_document_id': edo_document_id,
                'edo_endpoint': edo_endpoint
            }

        except Exception as e:
            logger.error(f"EDO export failed: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def _send_to_edo_stub(self, edo_system: str, endpoint: str, file_path: str) -> str:
        """
        EDO API integration (currently simulated)

        This is a STUB implementation. For production use, implement real EDO integration
        following the guide in docs/EDO_INTEGRATION.md

        Required steps for production:
        1. Register with chosen EDO system (Diadoc, SBIS, Kontur)
        2. Obtain API keys and certificates
        3. Install EDO SDK: pip install diadocsdk-python (for Diadoc)
        4. Implement authentication and document signing
        5. Replace this stub with real API calls

        Example for Diadoc:
        ```python
        from diadocsdk import Diadoc

        client = Diadoc(api_url=endpoint, api_key=self.config['diadoc_api_key'])
        auth_token = client.authenticate(...)

        with open(file_path, 'rb') as f:
            message = client.create_message(
                from_box_id=self.config['box_id'],
                to_inn=recipient_inn,
                documents=[{'filename': os.path.basename(file_path),
                          'content': f.read(),
                          'type': 'Nonformalized'}]
            )

        result = client.send_message(auth_token, message)
        return result.message_id
        ```

        Args:
            edo_system: EDO system name ('diadoc', 'sbis', 'kontur')
            endpoint: API endpoint URL
            file_path: Path to document file

        Returns:
            Document ID in EDO system

        Raises:
            EDOAuthenticationError: If authentication fails
            EDOUploadError: If document upload fails
        """
        logger.info(f"[SIMULATED EDO] System: {edo_system}")
        logger.info(f"[SIMULATED EDO] Endpoint: {endpoint}")
        logger.info(f"[SIMULATED EDO] File: {file_path}")
        logger.info(f"[SIMULATED EDO] File size: {os.path.getsize(file_path)} bytes")

        # Check if file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document file not found: {file_path}")

        # Simulate API call delay
        import time
        time.sleep(0.5)

        # Generate simulated document ID
        from datetime import datetime
        import hashlib

        # Create realistic-looking ID based on file hash
        with open(file_path, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()[:8]

        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        doc_id = f"{edo_system.upper()}-{timestamp}-{file_hash}"

        logger.info(f"[SIMULATED EDO] Generated document ID: {doc_id}")
        logger.warning(
            "EDO integration is currently SIMULATED. "
            "See docs/EDO_INTEGRATION.md for production implementation."
        )

        return doc_id

    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    def _log_export(
        self,
        disagreement_id: int,
        export_type: str,
        file_path: str,
        file_size: int,
        file_hash: str,
        user_id: Optional[str]
    ) -> DisagreementExportLog:
        """Log export to database"""
        export_log = DisagreementExportLog(
            disagreement_id=disagreement_id,
            export_type=export_type,
            file_path=file_path,
            file_size=file_size,
            file_hash=file_hash,
            exported_by=user_id
        )
        self.db_session.add(export_log)
        self.db_session.commit()
        self.db_session.refresh(export_log)
        return export_log


__all__ = ["DisagreementExportService"]

"""
Text Extraction Service
Извлекает текст из PDF, DOCX, TXT, XML, HTML и изображений (с OCR)
Сохраняет оригинальные bytes файлов и DOCX-версию для редактирования

Supports:
- PDF documents (native text extraction + конвертация в DOCX)
- DOCX files (сохранение оригинала)
- TXT plain text
- XML documents (contract XML formats)
- HTML documents
- Scanned images/PDFs (PaddleOCR)
"""

import time
import logging
from pathlib import Path
from typing import Dict, Any, Optional, BinaryIO, Union
from dataclasses import dataclass, field
import io

# PDF processing
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# DOCX processing
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# PDF to DOCX conversion
try:
    from pdf2docx import Converter as Pdf2DocxConverter
except ImportError:
    Pdf2DocxConverter = None

# OCR processing
try:
    from paddleocr import PaddleOCR
    from PIL import Image
    import pdf2image
except ImportError:
    PaddleOCR = None
    Image = None
    pdf2image = None

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """Результат извлечения текста"""
    text: str
    method: str  # 'pdf', 'docx', 'ocr'
    pages: int
    confidence: Optional[float]  # Для OCR
    processing_time: float  # секунды
    metadata: Dict[str, Any]  # Дополнительная информация
    # Новые поля для сохранения форматирования
    original_file_bytes: Optional[bytes] = None  # Бинарное содержимое оригинального файла
    docx_file_bytes: Optional[bytes] = None  # DOCX-версия документа (для редактирования)
    original_format: str = ''  # Исходный формат (pdf/docx/txt/xml/html)


class TextExtractor:
    """
    Сервис извлечения текста из документов

    Автоматически определяет тип документа и применяет
    оптимальный метод извлечения текста.
    Сохраняет оригинальные bytes и DOCX-версию для работы с форматированием.
    """

    def __init__(self, use_ocr: bool = True, ocr_lang: str = 'ru'):
        """
        Args:
            use_ocr: Использовать ли OCR для сканов
            ocr_lang: Язык для OCR ('ru', 'en', 'ru+en')
        """
        self.use_ocr = use_ocr
        self.ocr_lang = ocr_lang
        self._ocr_engine = None

        # Check dependencies
        if not pdfplumber:
            logger.warning("pdfplumber not installed. PDF extraction will be limited.")
        if not DocxDocument:
            logger.warning("python-docx not installed. DOCX extraction unavailable.")
        if not PaddleOCR and use_ocr:
            logger.warning("paddleocr not installed. OCR unavailable.")
        if not Pdf2DocxConverter:
            logger.warning("pdf2docx not installed. PDF→DOCX conversion unavailable.")

    @property
    def ocr_engine(self):
        """Lazy initialization of OCR engine"""
        if self._ocr_engine is None and PaddleOCR and self.use_ocr:
            try:
                self._ocr_engine = PaddleOCR(
                    use_angle_cls=True,
                    lang=self.ocr_lang,
                    show_log=False
                )
                logger.info(f"PaddleOCR initialized with lang={self.ocr_lang}")
            except Exception as e:
                logger.error(f"Failed to initialize PaddleOCR: {e}")
                self._ocr_engine = None
        return self._ocr_engine

    def _read_file_bytes(self, file_path: Union[Path, BinaryIO]) -> bytes:
        """Читает bytes из файла или file-like объекта"""
        if isinstance(file_path, Path):
            return file_path.read_bytes()
        else:
            pos = file_path.tell()
            data = file_path.read()
            file_path.seek(pos)
            return data if isinstance(data, bytes) else data.encode('utf-8')

    def _convert_pdf_to_docx(self, pdf_path: Union[Path, str], pdf_bytes: bytes) -> Optional[bytes]:
        """Конвертирует PDF в DOCX с сохранением форматирования"""
        if not Pdf2DocxConverter:
            logger.warning("pdf2docx not available, skipping PDF→DOCX conversion")
            return None

        import tempfile
        import os

        tmp_pdf = None
        tmp_docx = None
        try:
            # Если file_path — это BinaryIO, сохраняем во временный файл
            if not isinstance(pdf_path, (str, Path)):
                tmp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                tmp_pdf.write(pdf_bytes)
                tmp_pdf.close()
                pdf_path_str = tmp_pdf.name
            else:
                pdf_path_str = str(pdf_path)

            # Создаём временный файл для DOCX
            tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            tmp_docx.close()
            docx_path = tmp_docx.name

            # Конвертация
            cv = Pdf2DocxConverter(pdf_path_str)
            cv.convert(docx_path)
            cv.close()

            # Читаем результат
            with open(docx_path, 'rb') as f:
                docx_bytes = f.read()

            logger.info(f"PDF→DOCX conversion successful: {len(docx_bytes)} bytes")
            return docx_bytes

        except Exception as e:
            logger.error(f"PDF→DOCX conversion failed: {e}")
            return None
        finally:
            # Cleanup
            if tmp_pdf and os.path.exists(tmp_pdf.name):
                os.unlink(tmp_pdf.name)
            if tmp_docx and os.path.exists(tmp_docx.name):
                os.unlink(tmp_docx.name)

    def extract(self, file_path: Union[str, Path, BinaryIO],
                file_extension: Optional[str] = None) -> ExtractionResult:
        """
        Извлекает текст из документа

        Args:
            file_path: Путь к файлу или file-like объект
            file_extension: Расширение файла (если file_path - это BinaryIO)

        Returns:
            ExtractionResult с извлеченным текстом, метаданными и bytes файлов
        """
        start_time = time.time()

        # Определяем расширение файла
        if isinstance(file_path, (str, Path)):
            file_path = Path(file_path)
            ext = file_path.suffix.lower()
        else:
            ext = file_extension.lower() if file_extension else None

        # Выбираем метод извлечения
        if ext in ['.pdf']:
            result = self._extract_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            result = self._extract_from_docx(file_path)
        elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            result = self._extract_from_image(file_path)
        elif ext in ['.txt', '.text']:
            result = self._extract_from_txt(file_path)
        elif ext in ['.xml']:
            result = self._extract_from_xml(file_path)
        elif ext in ['.html', '.htm']:
            result = self._extract_from_html(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        # Добавляем время обработки
        result.processing_time = time.time() - start_time

        logger.info(f"Text extracted: method={result.method}, "
                   f"pages={result.pages}, "
                   f"chars={len(result.text)}, "
                   f"format={result.original_format}, "
                   f"has_docx={'yes' if result.docx_file_bytes else 'no'}, "
                   f"time={result.processing_time:.2f}s")

        return result

    def _extract_from_pdf(self, file_path: Union[Path, BinaryIO]) -> ExtractionResult:
        """
        Извлекает текст из PDF.
        Также конвертирует PDF → DOCX для сохранения форматирования.
        """
        if not pdfplumber:
            raise ImportError("pdfplumber is required for PDF extraction")

        try:
            # Читаем оригинальные bytes
            original_bytes = self._read_file_bytes(file_path)

            if isinstance(file_path, Path):
                pdf = pdfplumber.open(file_path)
            else:
                file_path.seek(0)
                pdf = pdfplumber.open(file_path)

            text_parts = []
            total_chars = 0

            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
                total_chars += len(page_text.strip())

            pdf.close()

            full_text = "\n\n".join(text_parts)
            pages = len(text_parts)

            # Если текста мало (< 100 символов на страницу) - это скан
            avg_chars_per_page = total_chars / pages if pages > 0 else 0

            if avg_chars_per_page < 100 and self.use_ocr:
                logger.info(f"Low text density ({avg_chars_per_page:.0f} chars/page). "
                           "Switching to OCR...")
                return self._extract_from_pdf_with_ocr(file_path)

            # Конвертируем PDF → DOCX для редактирования
            docx_bytes = self._convert_pdf_to_docx(file_path, original_bytes)

            return ExtractionResult(
                text=full_text,
                method='pdf',
                pages=pages,
                confidence=None,
                processing_time=0,
                metadata={
                    'chars_per_page': avg_chars_per_page,
                    'total_chars': total_chars
                },
                original_file_bytes=original_bytes,
                docx_file_bytes=docx_bytes,
                original_format='pdf'
            )

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise

    def _extract_from_pdf_with_ocr(self, file_path: Union[Path, BinaryIO]) -> ExtractionResult:
        """
        Извлекает текст из PDF с помощью OCR
        (для отсканированных документов)
        """
        if not pdf2image or not self.ocr_engine:
            raise ImportError("pdf2image and paddleocr required for PDF OCR")

        try:
            # Читаем оригинальные bytes
            original_bytes = self._read_file_bytes(file_path)

            # Конвертируем PDF в изображения
            if isinstance(file_path, Path):
                images = pdf2image.convert_from_path(file_path)
            else:
                file_path.seek(0)
                images = pdf2image.convert_from_bytes(file_path.read())

            text_parts = []
            total_confidence = 0
            conf_count = 0

            for i, image in enumerate(images):
                logger.debug(f"OCR processing page {i+1}/{len(images)}...")

                # OCR на изображение
                result = self.ocr_engine.ocr(image, cls=True)

                # Извлекаем текст и confidence
                page_text_parts = []
                for line in result[0] if result[0] else []:
                    text = line[1][0]  # Текст
                    conf = line[1][1]  # Confidence
                    page_text_parts.append(text)
                    total_confidence += conf
                    conf_count += 1

                page_text = "\n".join(page_text_parts)
                text_parts.append(page_text)

            full_text = "\n\n".join(text_parts)
            avg_confidence = total_confidence / conf_count if conf_count > 0 else 0.0

            # Конвертируем PDF → DOCX
            docx_bytes = self._convert_pdf_to_docx(file_path, original_bytes)

            return ExtractionResult(
                text=full_text,
                method='ocr',
                pages=len(images),
                confidence=avg_confidence,
                processing_time=0,
                metadata={
                    'ocr_lang': self.ocr_lang,
                    'lines_detected': conf_count
                },
                original_file_bytes=original_bytes,
                docx_file_bytes=docx_bytes,
                original_format='pdf'
            )

        except Exception as e:
            logger.error(f"PDF OCR extraction failed: {e}")
            raise

    def _extract_from_docx(self, file_path: Union[Path, BinaryIO]) -> ExtractionResult:
        """Извлекает текст из DOCX, сохраняя оригинальные bytes"""
        if not DocxDocument:
            raise ImportError("python-docx is required for DOCX extraction")

        try:
            # Читаем оригинальные bytes (это уже DOCX — используем и как docx_file_bytes)
            original_bytes = self._read_file_bytes(file_path)

            if isinstance(file_path, Path):
                doc = DocxDocument(file_path)
            else:
                file_path.seek(0)
                doc = DocxDocument(file_path)

            # Извлекаем текст из параграфов
            paragraphs = [p.text for p in doc.paragraphs]

            # Извлекаем текст из таблиц
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    table_texts.append(row_text)

            full_text = "\n".join(paragraphs)
            if table_texts:
                full_text += "\n\n" + "\n".join(table_texts)

            return ExtractionResult(
                text=full_text,
                method='docx',
                pages=1,
                confidence=None,
                processing_time=0,
                metadata={
                    'paragraphs': len(paragraphs),
                    'tables': len(doc.tables),
                    'total_chars': len(full_text)
                },
                original_file_bytes=original_bytes,
                docx_file_bytes=original_bytes,  # DOCX уже в нужном формате
                original_format='docx'
            )

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    def _extract_from_image(self, file_path: Union[Path, BinaryIO]) -> ExtractionResult:
        """Извлекает текст из изображения с помощью OCR"""
        if not Image or not self.ocr_engine:
            raise ImportError("PIL and paddleocr required for image OCR")

        try:
            original_bytes = self._read_file_bytes(file_path)

            # Открываем изображение
            if isinstance(file_path, Path):
                image = Image.open(file_path)
            else:
                file_path.seek(0)
                image = Image.open(file_path)

            # OCR
            result = self.ocr_engine.ocr(image, cls=True)

            # Извлекаем текст
            text_parts = []
            total_confidence = 0
            conf_count = 0

            for line in result[0] if result[0] else []:
                text = line[1][0]
                conf = line[1][1]
                text_parts.append(text)
                total_confidence += conf
                conf_count += 1

            full_text = "\n".join(text_parts)
            avg_confidence = total_confidence / conf_count if conf_count > 0 else 0.0

            return ExtractionResult(
                text=full_text,
                method='ocr',
                pages=1,
                confidence=avg_confidence,
                processing_time=0,
                metadata={
                    'ocr_lang': self.ocr_lang,
                    'lines_detected': conf_count,
                    'image_size': image.size
                },
                original_file_bytes=original_bytes,
                docx_file_bytes=None,
                original_format='image'
            )

        except Exception as e:
            logger.error(f"Image OCR extraction failed: {e}")
            raise

    def _extract_from_txt(self, file_path: Union[Path, BinaryIO]) -> ExtractionResult:
        """Извлекает текст из TXT файла"""
        try:
            original_bytes = self._read_file_bytes(file_path)

            if isinstance(file_path, (str, Path)):
                file_path = Path(file_path)
                text = file_path.read_text(encoding='utf-8')
            else:
                text = original_bytes.decode('utf-8') if isinstance(original_bytes, bytes) else original_bytes

            return ExtractionResult(
                text=text,
                method='txt',
                pages=1,
                confidence=None,
                processing_time=0,
                metadata={
                    'total_chars': len(text),
                    'lines': text.count('\n') + 1
                },
                original_file_bytes=original_bytes,
                docx_file_bytes=None,
                original_format='txt'
            )

        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise

    def _extract_from_xml(self, file_path: Union[Path, BinaryIO]) -> ExtractionResult:
        """Извлекает текст из XML файла (убирает теги, сохраняет текстовое содержимое)"""
        import xml.etree.ElementTree as ET

        try:
            original_bytes = self._read_file_bytes(file_path)

            if isinstance(file_path, (str, Path)):
                file_path = Path(file_path)
                raw_content = file_path.read_text(encoding='utf-8')
            else:
                raw_content = original_bytes.decode('utf-8') if isinstance(original_bytes, bytes) else original_bytes

            # Парсим XML и извлекаем весь текст
            root = ET.fromstring(raw_content)

            def extract_text_recursive(element):
                """Рекурсивно извлекает текст из XML элементов"""
                parts = []
                if element.text and element.text.strip():
                    parts.append(element.text.strip())
                for child in element:
                    parts.extend(extract_text_recursive(child))
                    if child.tail and child.tail.strip():
                        parts.append(child.tail.strip())
                return parts

            text_parts = extract_text_recursive(root)
            text = "\n".join(text_parts)

            return ExtractionResult(
                text=text,
                method='xml',
                pages=1,
                confidence=None,
                processing_time=0,
                metadata={
                    'total_chars': len(text),
                    'lines': text.count('\n') + 1,
                    'root_tag': root.tag,
                    'raw_xml_chars': len(raw_content)
                },
                original_file_bytes=original_bytes,
                docx_file_bytes=None,
                original_format='xml'
            )

        except ET.ParseError as e:
            logger.warning(f"XML parse error, falling back to text extraction: {e}")
            # Fallback: если XML невалидный, читаем как текст
            if isinstance(file_path, Path):
                text = file_path.read_text(encoding='utf-8')
            else:
                file_path.seek(0)
                raw = file_path.read()
                text = raw.decode('utf-8') if isinstance(raw, bytes) else raw

            # Простое удаление тегов через regex
            import re
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()

            return ExtractionResult(
                text=text,
                method='xml_fallback',
                pages=1,
                confidence=None,
                processing_time=0,
                metadata={'total_chars': len(text), 'parse_error': str(e)},
                original_file_bytes=original_bytes,
                docx_file_bytes=None,
                original_format='xml'
            )

        except Exception as e:
            logger.error(f"XML extraction failed: {e}")
            raise

    def _extract_from_html(self, file_path: Union[Path, BinaryIO]) -> ExtractionResult:
        """Извлекает текст из HTML файла (убирает теги, сохраняет текст)"""
        import re

        try:
            original_bytes = self._read_file_bytes(file_path)

            if isinstance(file_path, (str, Path)):
                file_path = Path(file_path)
                raw_content = file_path.read_text(encoding='utf-8')
            else:
                raw_content = original_bytes.decode('utf-8') if isinstance(original_bytes, bytes) else original_bytes

            # Удаляем script и style блоки
            text = re.sub(r'<script[^>]*>.*?</script>', '', raw_content, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)

            # Заменяем блочные теги на переносы строк
            text = re.sub(r'<br\s*/?>|</?p>|</?div>|</?tr>|</?li>', '\n', text, flags=re.IGNORECASE)

            # Убираем все оставшиеся теги
            text = re.sub(r'<[^>]+>', ' ', text)

            # Декодируем HTML entities
            import html
            text = html.unescape(text)

            # Убираем лишние пробелы
            text = re.sub(r'[ \t]+', ' ', text)
            text = re.sub(r'\n\s*\n', '\n\n', text)
            text = text.strip()

            return ExtractionResult(
                text=text,
                method='html',
                pages=1,
                confidence=None,
                processing_time=0,
                metadata={
                    'total_chars': len(text),
                    'lines': text.count('\n') + 1,
                    'raw_html_chars': len(raw_content)
                },
                original_file_bytes=original_bytes,
                docx_file_bytes=None,
                original_format='html'
            )

        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            raise

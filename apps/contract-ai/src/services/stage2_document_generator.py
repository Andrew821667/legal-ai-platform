# -*- coding: utf-8 -*-
"""
Stage 2 Final Document Generator (2.4)

Provides two output variants:
- Variant A: Corrected DOCX with accepted recommendations
- Variant B: Disagreement protocol DOCX/JSON (before -> after)
"""

from datetime import datetime
from io import BytesIO
import json
from typing import Any, Dict, List, Optional

from docx import Document


class Stage2DocumentGenerator:
    """Generate final outputs for Stage 2."""

    def generate_corrected_docx(
        self,
        base_docx_bytes: Optional[bytes],
        accepted_recommendations: List[Dict[str, Any]],
        source_file_name: str = "contract",
    ) -> bytes:
        """
        Variant A:
        Build corrected DOCX based on extracted DOCX and append structured changes.
        """

        if base_docx_bytes:
            document = Document(BytesIO(base_docx_bytes))
        else:
            document = Document()
            document.add_heading("Договор (восстановленная версия)", level=1)

        document.add_page_break()
        document.add_heading("Внесенные правки (Stage 2.4)", level=1)
        document.add_paragraph(f"Документ-источник: {source_file_name}")
        document.add_paragraph(f"Дата генерации: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

        if not accepted_recommendations:
            document.add_paragraph("Принятые рекомендации отсутствуют.")
            return self._to_bytes(document)

        document.add_paragraph(
            "Ниже перечислены согласованные изменения, которые должны быть учтены в рабочей редакции договора."
        )

        for idx, rec in enumerate(accepted_recommendations, 1):
            section = self._build_section_label(rec)
            action_type = rec.get("action_type", "modify")
            priority = rec.get("priority", "optional")
            reason = rec.get("reason", "")
            source = rec.get("source", "section_analysis")
            original_text = rec.get("original_text", "")
            proposed_text = rec.get("proposed_text", "")

            document.add_heading(f"{idx}. {section}", level=2)
            document.add_paragraph(f"Тип правки: {self._action_label(action_type)}")
            document.add_paragraph(f"Приоритет: {priority}")
            document.add_paragraph(f"Источник: {source}")

            if reason:
                document.add_paragraph(f"Обоснование: {reason}")

            if original_text:
                p_old = document.add_paragraph()
                p_old.add_run("Было: ").bold = True
                p_old.add_run(original_text)

            p_new = document.add_paragraph()
            p_new.add_run("Стало: ").bold = True
            p_new.add_run(proposed_text or "Текст правки не указан")

        return self._to_bytes(document)

    def generate_disagreement_protocol_docx(
        self,
        accepted_recommendations: List[Dict[str, Any]],
        source_file_name: str = "contract",
    ) -> bytes:
        """
        Variant B:
        Generate disagreement protocol DOCX (before -> after table).
        """
        document = Document()
        document.add_heading("Протокол разногласий", level=1)
        document.add_paragraph(f"Документ: {source_file_name}")
        document.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

        if not accepted_recommendations:
            document.add_paragraph("Согласованные разногласия отсутствуют.")
            return self._to_bytes(document)

        table = document.add_table(rows=1, cols=6)
        header_cells = table.rows[0].cells
        header_cells[0].text = "№"
        header_cells[1].text = "Раздел"
        header_cells[2].text = "Тип"
        header_cells[3].text = "Было"
        header_cells[4].text = "Стало"
        header_cells[5].text = "Обоснование"

        for idx, rec in enumerate(accepted_recommendations, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = self._build_section_label(rec)
            row[2].text = self._action_label(rec.get("action_type", "modify"))
            row[3].text = str(rec.get("original_text", ""))
            row[4].text = str(rec.get("proposed_text", ""))
            row[5].text = str(rec.get("reason", ""))

        return self._to_bytes(document)

    def generate_disagreement_protocol_json(
        self,
        accepted_recommendations: List[Dict[str, Any]],
    ) -> str:
        """JSON export for protocol."""
        payload: List[Dict[str, Any]] = []
        for idx, rec in enumerate(accepted_recommendations, 1):
            payload.append(
                {
                    "index": idx,
                    "section": self._build_section_label(rec),
                    "action_type": rec.get("action_type", "modify"),
                    "priority": rec.get("priority", "optional"),
                    "before": rec.get("original_text", ""),
                    "after": rec.get("proposed_text", ""),
                    "reason": rec.get("reason", ""),
                    "source": rec.get("source", "section_analysis"),
                }
            )

        return json.dumps(payload, ensure_ascii=False, indent=2)

    def _to_bytes(self, document: Document) -> bytes:
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _build_section_label(self, rec: Dict[str, Any]) -> str:
        section_number = rec.get("section_number", "")
        section_title = rec.get("section_title", "")
        if section_number and section_title:
            return "{0}. {1}".format(section_number, section_title)
        if section_title:
            return str(section_title)
        if section_number:
            return str(section_number)
        return "Не указан"

    def _action_label(self, action_type: str) -> str:
        mapping = {
            "add": "Добавить",
            "modify": "Изменить",
            "remove": "Удалить",
            "missing": "Добавить отсутствующий пункт",
            "weakened": "Усилить условие",
            "contradicts": "Исправить противоречие",
            "added": "Проверить добавленное условие",
        }
        return mapping.get(str(action_type).lower(), "Изменить")


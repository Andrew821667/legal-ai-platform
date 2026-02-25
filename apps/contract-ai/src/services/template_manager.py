# -*- coding: utf-8 -*-
"""
Template Manager - system for managing contract templates

Features:
- Loading and storing templates
- Versioning
- Variable substitution (Jinja2)
- Structure validation
- Template inheritance
- Export to DOCX
"""
import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
from lxml import etree
from jinja2 import Template as Jinja2Template, Environment, BaseLoader, TemplateNotFound
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger
from sqlalchemy.orm import Session

from ..models.database import Template as TemplateModel
from ..models.repositories import TemplateRepository


class TemplateManager:
    """
    Contract template manager

    Functionality:
    - Load XML templates
    - Versioning
    - Variable substitution via Jinja2
    - Structure validation
    - Inheritance from base templates
    - Convert XML to DOCX
    """

    def __init__(self, db_session: Session, templates_dir: str = "data/templates"):
        """
        Args:
            db_session: DB session for repository
            templates_dir: Directory with templates on disk
        """
        self.db = db_session
        self.repository = TemplateRepository(db_session)
        self.templates_dir = Path(templates_dir)
        self.templates_dir.mkdir(parents=True, exist_ok=True)

        # Define required elements in template
        self.required_elements = ['metadata', 'parties', 'clauses']

        logger.info(f"TemplateManager initialized with templates_dir: {self.templates_dir}")

    def load_template_from_file(
        self,
        file_path: str,
        contract_type: str,
        name: str,
        version: str = "1.0",
        base_template_id: Optional[str] = None,
        created_by: Optional[str] = None
    ) -> TemplateModel:
        """
        Load template from XML file into DB

        Args:
            file_path: Path to XML template file
            contract_type: Contract type (supply, service, lease, etc.)
            name: Template name
            version: Template version
            base_template_id: Base template ID (for inheritance)
            created_by: Creator user ID

        Returns:
            Created Template object

        Raises:
            FileNotFoundError: If file not found
            ValueError: If template is not valid
        """
        logger.info(f"Loading template from file: {file_path}")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Template file not found: {file_path}")

        # Read XML
        with open(file_path, 'r', encoding='utf-8') as f:
            xml_content = f.read()

        # Validate
        if not self.validate_template(xml_content):
            raise ValueError(f"Template validation failed for: {file_path}")

        # If there is a base template, apply inheritance
        if base_template_id:
            xml_content = self._apply_inheritance(xml_content, base_template_id)

        # Extract template structure
        structure = self._extract_template_structure(xml_content)

        # Create DB record
        template = self.repository.create(
            name=name,
            contract_type=contract_type,
            xml_content=xml_content,
            structure=structure,
            version=version,
            created_by=created_by,
            active=True
        )

        logger.info(f"Template loaded successfully: {template.id} ({name} v{version})")
        return template

    def get_template(
        self,
        contract_type: str,
        version: Optional[str] = None
    ) -> Optional[TemplateModel]:
        """
        Get template from DB

        Args:
            contract_type: Contract type
            version: Version (if None, returns latest)

        Returns:
            Template or None
        """
        if version:
            templates = self.repository.get_by_type(contract_type, active_only=True)
            for template in templates:
                if template.version == version:
                    return template
            return None
        else:
            return self.repository.get_latest_version(contract_type)

    def fill_template(
        self,
        template_id: str,
        data: Dict[str, Any],
        validate_output: bool = True
    ) -> str:
        """
        Fill template with data using Jinja2

        Args:
            template_id: Template ID
            data: Dictionary with data for substitution
            validate_output: Whether to validate result

        Returns:
            Filled XML

        Raises:
            ValueError: If template not found or result is not valid
        """
        logger.info(f"Filling template: {template_id}")

        template = self.repository.get_by_id(template_id)
        if not template:
            raise ValueError(f"Template not found: {template_id}")

        # Use Jinja2 for variable substitution
        jinja_template = Jinja2Template(template.xml_content)
        filled_xml = jinja_template.render(**data)

        # Validate result
        if validate_output and not self.validate_xml(filled_xml):
            raise ValueError("Filled template is not valid XML")

        logger.info(f"Template filled successfully")
        return filled_xml

    def validate_template(self, xml_content: str) -> bool:
        """
        Validate XML template structure

        Args:
            xml_content: XML string

        Returns:
            True if valid
        """
        try:
            root = etree.fromstring(xml_content.encode('utf-8'))

            # Check root element
            if root.tag != 'contract':
                logger.error(f"Invalid root element: {root.tag}, expected 'contract'")
                return False

            # Check required elements
            for element in self.required_elements:
                if root.find(element) is None:
                    logger.error(f"Missing required element: {element}")
                    return False

            logger.info("Template validation successful")
            return True

        except etree.XMLSyntaxError as e:
            logger.error(f"XML syntax error: {e}")
            return False

    def validate_xml(self, xml_content: str) -> bool:
        """
        Validate XML (general validation)

        Args:
            xml_content: XML string

        Returns:
            True if valid
        """
        try:
            etree.fromstring(xml_content.encode('utf-8'))
            return True
        except etree.XMLSyntaxError as e:
            logger.error(f"XML validation failed: {e}")
            return False

    def create_version(
        self,
        base_template_id: str,
        new_version: str,
        changes: Optional[Dict[str, Any]] = None
    ) -> TemplateModel:
        """
        Create new version of existing template

        Args:
            base_template_id: Base template ID
            new_version: New version number
            changes: Changes to apply (optional)

        Returns:
            New Template

        Raises:
            ValueError: If base template not found
        """
        logger.info(f"Creating new version {new_version} from template {base_template_id}")

        base_template = self.repository.get_by_id(base_template_id)
        if not base_template:
            raise ValueError(f"Base template not found: {base_template_id}")

        # Copy content
        new_xml_content = base_template.xml_content

        # Apply changes if any
        if changes:
            new_xml_content = self._apply_changes(new_xml_content, changes)

        # Create new version
        new_template = self.repository.create(
            name=base_template.name,
            contract_type=base_template.contract_type,
            xml_content=new_xml_content,
            structure=self._extract_template_structure(new_xml_content),
            version=new_version,
            created_by=base_template.created_by,
            active=True
        )

        logger.info(f"New version created: {new_template.id}")
        return new_template

    def export_to_docx(
        self,
        filled_xml: str,
        output_path: str,
        style_config: Optional[Dict[str, Any]] = None
    ) -> None:
        """
        Convert filled XML to DOCX file

        Args:
            filled_xml: Filled XML
            output_path: Path to save DOCX
            style_config: Style settings (fonts, spacing, etc.)
        """
        logger.info(f"Exporting to DOCX: {output_path}")

        # Parse XML
        root = etree.fromstring(filled_xml.encode('utf-8'))

        # Create new DOCX document
        doc = Document()

        # Apply styles if specified
        if style_config is None:
            style_config = self._get_default_style_config()

        # Extract metadata
        metadata = root.find('metadata')
        if metadata is not None:
            title = metadata.findtext('title', 'Contract')
            doc.core_properties.title = title

        # Add title
        title_elem = root.find('.//metadata/title')
        if title_elem is not None and title_elem.text:
            heading = doc.add_heading(title_elem.text, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add preamble (if any)
        preamble = root.find('.//clause[@type="preamble"]')
        if preamble is not None:
            content = preamble.find('content')
            if content is not None:
                for para in content.findall('paragraph'):
                    if para.text:
                        p = doc.add_paragraph(para.text)
                        self._apply_paragraph_style(p, style_config)

        # Add sections
        clauses = root.find('clauses')
        if clauses is not None:
            for clause in clauses.findall('clause'):
                if clause.get('type') == 'preamble':
                    continue  # Already added

                # Section title
                title = clause.findtext('title', 'Section')
                doc.add_heading(title, level=1)

                # Section content
                content = clause.find('content')
                if content is not None:
                    for para in content.findall('paragraph'):
                        if para.text:
                            p = doc.add_paragraph(para.text)
                            self._apply_paragraph_style(p, style_config)

        # Add tables
        tables = root.find('tables')
        if tables is not None:
            for table_elem in tables.findall('table'):
                rows = table_elem.findall('row')
                if not rows:
                    continue

                # Determine table size
                num_rows = len(rows)
                num_cols = len(rows[0].findall('cell')) if rows else 0

                if num_cols == 0:
                    continue

                # Create table
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'

                # Fill table
                for i, row in enumerate(rows):
                    cells = row.findall('cell')
                    for j, cell in enumerate(cells):
                        if cell.text:
                            table.rows[i].cells[j].text = cell.text

        # Add signatures (if any)
        parties = root.find('parties')
        if parties is not None:
            doc.add_paragraph()  # Empty line
            doc.add_paragraph('Signatures:')

            for party in parties.findall('party'):
                role = party.get('role', 'unknown')
                name = party.findtext('name', 'Not specified')

                role_name = {
                    'supplier': 'Supplier',
                    'buyer': 'Buyer',
                    'seller': 'Seller',
                    'customer': 'Customer',
                    'contractor': 'Contractor'
                }.get(role, role.capitalize())

                p = doc.add_paragraph()
                p.add_run(f'{role_name}: ').bold = True
                p.add_run(f'{name}')
                p.add_run('\n\n__________________ / ________________')

        # Save
        doc.save(output_path)
        logger.info(f"DOCX exported successfully: {output_path}")

    def get_template_variables(self, template_id: str) -> List[str]:
        """
        Extract list of variables from template

        Args:
            template_id: Template ID

        Returns:
            List of variable names (e.g., ['company_name', 'amount'])
        """
        template = self.repository.get_by_id(template_id)
        if not template:
            return []

        # Find Jinja2 variables like {{ variable_name }}
        pattern = r'\{\{\s*(\w+)\s*\}\}'
        variables = re.findall(pattern, template.xml_content)

        # Remove duplicates and sort
        return sorted(list(set(variables)))

    def list_templates(
        self,
        contract_type: Optional[str] = None,
        active_only: bool = True
    ) -> List[TemplateModel]:
        """
        Return list of templates

        Args:
            contract_type: Filter by type (if None, all types)
            active_only: Only active templates

        Returns:
            List of Templates
        """
        if contract_type:
            return self.repository.get_by_type(contract_type, active_only)
        else:
            return self.repository.get_active_templates() if active_only else self.repository.get_all()

    def deactivate_template(self, template_id: str) -> bool:
        """
        Deactivate template

        Args:
            template_id: Template ID

        Returns:
            True if successful
        """
        template = self.repository.update(template_id, active=False)
        return template is not None

    # --- Helper methods ---

    def _extract_template_structure(self, xml_content: str) -> str:
        """
        Extract template structure (for DB storage)

        Returns:
            JSON string with structure
        """
        import json

        try:
            root = etree.fromstring(xml_content.encode('utf-8'))

            structure = {
                'sections': [],
                'has_tables': False,
                'has_parties': False
            }

            # Sections
            clauses = root.find('clauses')
            if clauses is not None:
                for clause in clauses.findall('clause'):
                    structure['sections'].append({
                        'id': clause.get('id'),
                        'type': clause.get('type'),
                        'title': clause.findtext('title', '')
                    })

            # Tables
            tables = root.find('tables')
            if tables is not None:
                structure['has_tables'] = len(tables.findall('table')) > 0

            # Parties
            parties = root.find('parties')
            if parties is not None:
                structure['has_parties'] = len(parties.findall('party')) > 0

            return json.dumps(structure, ensure_ascii=False)

        except Exception as e:
            logger.error(f"Error extracting structure: {e}")
            return "{}"

    def _apply_inheritance(self, child_xml: str, base_template_id: str) -> str:
        """
        Apply template inheritance
        Merge child template with base

        Args:
            child_xml: Child template XML
            base_template_id: Base template ID

        Returns:
            Merged XML
        """
        logger.info(f"Applying inheritance from base template: {base_template_id}")

        base_template = self.repository.get_by_id(base_template_id)
        if not base_template:
            logger.warning(f"Base template not found: {base_template_id}, returning child as is")
            return child_xml

        # Parse both templates
        base_root = etree.fromstring(base_template.xml_content.encode('utf-8'))
        child_root = etree.fromstring(child_xml.encode('utf-8'))

        # Inheritance strategy:
        # 1. Take base template as foundation
        # 2. Override elements from child template

        result_root = etree.Element("contract")

        # Copy everything from base
        for element in base_root:
            result_root.append(element)

        # Override with elements from child
        for element in child_root:
            # Remove old element with same tag if exists
            existing = result_root.find(element.tag)
            if existing is not None:
                result_root.remove(existing)

            # Add new
            result_root.append(element)

        # Convert back to string
        result_xml = etree.tostring(result_root, encoding='unicode', pretty_print=True)
        result_xml = '<?xml version="1.0" encoding="UTF-8"?>\n' + result_xml

        return result_xml

    def _apply_changes(self, xml_content: str, changes: Dict[str, Any]) -> str:
        """
        Apply changes to XML template

        Args:
            xml_content: Source XML
            changes: Dictionary of changes

        Returns:
            Modified XML
        """
        # Simple implementation: string replacement
        # Can be extended for more complex changes

        result = xml_content

        for key, value in changes.items():
            # Replace variables
            result = result.replace(f'{{{{{key}}}}}', str(value))

        return result

    def _get_default_style_config(self) -> Dict[str, Any]:
        """
        Return default style configuration for DOCX
        """
        return {
            'font_name': 'Times New Roman',
            'font_size': 12,
            'line_spacing': 1.5,
            'paragraph_spacing': 6
        }

    def _apply_paragraph_style(self, paragraph, style_config: Dict[str, Any]) -> None:
        """
        Apply style to paragraph
        """
        for run in paragraph.runs:
            run.font.name = style_config.get('font_name', 'Times New Roman')
            run.font.size = Pt(style_config.get('font_size', 12))

        paragraph.paragraph_format.line_spacing = style_config.get('line_spacing', 1.5)
        paragraph.paragraph_format.space_after = Pt(style_config.get('paragraph_spacing', 6))


# Export
__all__ = ["TemplateManager"]
